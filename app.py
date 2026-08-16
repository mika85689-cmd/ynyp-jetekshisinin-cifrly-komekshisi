import os, io, csv
from datetime import date
from flask import Flask, request, jsonify, send_from_directory, session
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash

app = Flask(__name__, static_folder=None)
app.secret_key = os.environ.get('SECRET_KEY', 'dev-only-change-me')
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'sqlite:///synyp_ai.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_SECURE'] = os.environ.get('RENDER') == 'true'
db = SQLAlchemy(app)

SUBJECTS = ['Қазақ тілі','Қазақ әдебиеті','Орыс тілі','Ағылшын тілі','Математика','Алгебра','Геометрия','Информатика','Физика','Химия','Биология','География','Қазақстан тарихы','Дүниежүзі тарихы','Дене шынықтыру','Көркем еңбек','Музыка']

class User(db.Model):
    id=db.Column(db.Integer,primary_key=True); name=db.Column(db.String(160),nullable=False); login=db.Column(db.String(80),unique=True,nullable=False); password_hash=db.Column(db.String(255),nullable=False); role=db.Column(db.String(20),nullable=False,default='teacher'); class_name=db.Column(db.String(40),default=''); must_change=db.Column(db.Boolean,default=True)
class Student(db.Model):
    id=db.Column(db.Integer,primary_key=True); class_name=db.Column(db.String(40),nullable=False,index=True); name=db.Column(db.String(180),nullable=False); gender=db.Column(db.String(20),default=''); social=db.Column(db.String(120),default=''); address=db.Column(db.String(220),default=''); status=db.Column(db.String(40),default='Белсенді')
class Grade(db.Model):
    id=db.Column(db.Integer,primary_key=True); class_name=db.Column(db.String(40),nullable=False,index=True); student_id=db.Column(db.Integer,db.ForeignKey('student.id',ondelete='CASCADE'),nullable=False); subject=db.Column(db.String(120),nullable=False); term=db.Column(db.Integer,nullable=False); points=db.Column(db.Float,nullable=False); student=db.relationship('Student')
class Attendance(db.Model):
    id=db.Column(db.Integer,primary_key=True); class_name=db.Column(db.String(40),nullable=False,index=True); student_id=db.Column(db.Integer,db.ForeignKey('student.id',ondelete='CASCADE'),nullable=False); date=db.Column(db.String(20),nullable=False); status=db.Column(db.String(40),nullable=False); student=db.relationship('Student')
class Schedule(db.Model):
    id=db.Column(db.Integer,primary_key=True); class_name=db.Column(db.String(40),nullable=False,index=True); day=db.Column(db.String(30),nullable=False); lesson=db.Column(db.String(120),nullable=False); time=db.Column(db.String(20),nullable=False)
class Plan(db.Model):
    id=db.Column(db.Integer,primary_key=True); class_name=db.Column(db.String(40),nullable=False,index=True); type=db.Column(db.String(40),nullable=False); title=db.Column(db.String(220),nullable=False); date=db.Column(db.String(20),default=''); note=db.Column(db.Text,default='')

def user_json(u): return {'id':u.id,'name':u.name,'login':u.login,'role':u.role,'class_name':u.class_name,'must_change':u.must_change}
def current_user():
    uid=session.get('uid'); return db.session.get(User,uid) if uid else None
def require(admin=False):
    u=current_user()
    if not u: return None,(jsonify(error='Кіру қажет'),401)
    if admin and u.role!='admin': return None,(jsonify(error='Рұқсат жоқ'),403)
    return u,None
def scoped(model,u): return model.query if u.role=='admin' else model.query.filter_by(class_name=u.class_name)
def allowed_student(u,s): return s and (u.role=='admin' or s.class_name==u.class_name)

@app.route('/')
def root(): return send_from_directory('.', 'index.html')
@app.get('/health')
def health(): return jsonify(ok=True)
@app.post('/api/login')
def login():
    d=request.get_json(silent=True) or {}; u=User.query.filter_by(login=str(d.get('login','')).strip()).first()
    if not u or not check_password_hash(u.password_hash,str(d.get('password',''))): return jsonify(error='Логин немесе құпиясөз қате'),401
    session.clear(); session['uid']=u.id; return jsonify(ok=True,user=user_json(u))
@app.post('/api/logout')
def logout(): session.clear(); return jsonify(ok=True)
@app.get('/api/me')
def me():
    u,e=require(); return e or jsonify(user_json(u))
@app.post('/api/change-password')
def change_password():
    u,e=require();
    if e:return e
    p=str((request.get_json(silent=True) or {}).get('password',''))
    if len(p)<8:return jsonify(error='Құпиясөз кемінде 8 таңба болсын'),400
    u.password_hash=generate_password_hash(p); u.must_change=False; db.session.commit(); return jsonify(ok=True)
@app.get('/api/subjects')
def subjects():
    u,e=require(); return e or jsonify(SUBJECTS)

@app.route('/api/users',methods=['GET','POST'])
def users():
    u,e=require(admin=True)
    if e:return e
    if request.method=='GET': return jsonify([user_json(x) for x in User.query.order_by(User.id).all()])
    d=request.get_json(silent=True) or {}
    for k in ['name','login','class_name','password']:
        if not str(d.get(k,'')).strip(): return jsonify(error='Барлық өрісті толтырыңыз'),400
    if User.query.filter_by(login=d['login'].strip()).first(): return jsonify(error='Бұл логин бос емес'),409
    x=User(name=d['name'].strip(),login=d['login'].strip(),password_hash=generate_password_hash(d['password']),role='teacher',class_name=d['class_name'].strip(),must_change=True); db.session.add(x); db.session.commit(); return jsonify(id=x.id),201

@app.route('/api/students',methods=['GET','POST'])
def students():
    u,e=require();
    if e:return e
    if request.method=='GET':
        return jsonify([{'id':x.id,'class_name':x.class_name,'name':x.name,'gender':x.gender,'social':x.social,'address':x.address,'status':x.status} for x in scoped(Student,u).order_by(Student.class_name,Student.name).all()])
    d=request.get_json(silent=True) or {}; cls=str(d.get('class_name','')).strip() if u.role=='admin' else u.class_name
    if not cls or not str(d.get('name','')).strip():return jsonify(error='Сынып пен аты-жөні міндетті'),400
    x=Student(class_name=cls,name=d['name'].strip(),gender=d.get('gender',''),social=d.get('social',''),address=d.get('address',''),status=d.get('status','Белсенді')); db.session.add(x); db.session.commit(); return jsonify(id=x.id),201

@app.route('/api/grades',methods=['GET','POST'])
def grades():
    u,e=require();
    if e:return e
    if request.method=='GET':
        return jsonify([{'id':x.id,'class_name':x.class_name,'student_id':x.student_id,'student':x.student.name,'subject':x.subject,'term':x.term,'points':x.points} for x in scoped(Grade,u).order_by(Grade.id.desc()).all()])
    d=request.get_json(silent=True) or {}; s=db.session.get(Student,int(d.get('student_id',0) or 0))
    if not allowed_student(u,s):return jsonify(error='Оқушы табылмады'),404
    try:p=float(d.get('points'))
    except:return jsonify(error='Балл енгізіңіз'),400
    if p<0 or p>100:return jsonify(error='Балл 0–100 аралығында болуы керек'),400
    x=Grade(class_name=s.class_name,student_id=s.id,subject=str(d.get('subject','')).strip(),term=int(d.get('term',1)),points=p); db.session.add(x); db.session.commit(); return jsonify(ok=True),201

@app.route('/api/attendance',methods=['GET','POST'])
def attendance():
    u,e=require();
    if e:return e
    if request.method=='GET':
        return jsonify([{'id':x.id,'class_name':x.class_name,'student_id':x.student_id,'student':x.student.name,'date':x.date,'status':x.status} for x in scoped(Attendance,u).order_by(Attendance.id.desc()).all()])
    d=request.get_json(silent=True) or {}; s=db.session.get(Student,int(d.get('student_id',0) or 0))
    if not allowed_student(u,s):return jsonify(error='Оқушы табылмады'),404
    x=Attendance(class_name=s.class_name,student_id=s.id,date=d.get('date') or date.today().isoformat(),status=d.get('status','Қатысты')); db.session.add(x); db.session.commit(); return jsonify(ok=True),201

@app.route('/api/schedule',methods=['GET','POST'])
def schedule():
    u,e=require();
    if e:return e
    if request.method=='GET': return jsonify([{'id':x.id,'class_name':x.class_name,'day':x.day,'lesson':x.lesson,'time':x.time} for x in scoped(Schedule,u).order_by(Schedule.id.desc()).all()])
    d=request.get_json(silent=True) or {}; cls=str(d.get('class_name','')).strip() if u.role=='admin' else u.class_name
    x=Schedule(class_name=cls,day=d.get('day',''),lesson=d.get('lesson',''),time=d.get('time','')); db.session.add(x); db.session.commit(); return jsonify(ok=True),201

@app.route('/api/plans',methods=['GET','POST'])
def plans():
    u,e=require();
    if e:return e
    if request.method=='GET': return jsonify([{'id':x.id,'class_name':x.class_name,'type':x.type,'title':x.title,'date':x.date,'note':x.note} for x in scoped(Plan,u).order_by(Plan.id.desc()).all()])
    d=request.get_json(silent=True) or {}; cls=str(d.get('class_name','')).strip() if u.role=='admin' else u.class_name
    x=Plan(class_name=cls,type=d.get('type','Жоспар'),title=d.get('title',''),date=d.get('date',''),note=d.get('note','')); db.session.add(x); db.session.commit(); return jsonify(ok=True),201

@app.get('/api/dashboard')
def dashboard():
    u,e=require();
    if e:return e
    st=scoped(Student,u).count(); gs=scoped(Grade,u).all(); av=round(sum(x.points for x in gs)/len(gs),1) if gs else 0; ab=sum(1 for x in scoped(Attendance,u).all() if x.status!='Қатысты'); cl=User.query.filter_by(role='teacher').with_entities(User.class_name).distinct().count() if u.role=='admin' else 1
    return jsonify(students=st,avg=av,absent=ab,classes=cl)

@app.post('/api/import-students')
def import_students():
    u,e=require();
    if e:return e
    f=request.files.get('file'); cls=(request.form.get('class_name','').strip() if u.role=='admin' else u.class_name)
    if not f or not cls:return jsonify(error='Файл және сынып қажет'),400
    ext=os.path.splitext(f.filename or '')[1].lower(); rows=[]
    try:
        if ext=='.csv':
            text=f.read().decode('utf-8-sig'); vals=list(csv.reader(io.StringIO(text))); rows=vals[1:] if vals and any(('аты' in str(x).lower() or 'name' in str(x).lower()) for x in vals[0]) else vals
        elif ext=='.xlsx':
            from openpyxl import load_workbook
            vals=list(load_workbook(f.stream,read_only=True,data_only=True).active.iter_rows(values_only=True)); rows=vals[1:] if vals else []
        elif ext=='.docx':
            from docx import Document
            doc=Document(f.stream); vals=[[c.text.strip() for c in r.cells] for r in doc.tables[0].rows] if doc.tables else [[p.text.strip()] for p in doc.paragraphs if p.text.strip()]; rows=vals[1:] if vals else []
        else:return jsonify(error='CSV, XLSX немесе DOCX қолданыңыз'),400
    except Exception as ex:return jsonify(error='Файл оқылмады: '+str(ex)),400
    count=0
    for r in rows:
        r=list(r)+['','','']; name=str(r[0] or '').strip()
        if name: db.session.add(Student(class_name=cls,name=name,gender=str(r[1] or ''),social=str(r[2] or ''),address=str(r[3] or ''),status='Белсенді')); count+=1
    db.session.commit(); return jsonify(ok=True,count=count)

@app.delete('/api/<kind>/<int:item_id>')
def delete(kind,item_id):
    u,e=require();
    if e:return e
    models={'students':Student,'grades':Grade,'attendance':Attendance,'schedule':Schedule,'plans':Plan,'users':User}; m=models.get(kind)
    if not m:return jsonify(error='Қате бөлім'),404
    x=db.session.get(m,item_id)
    if not x:return jsonify(error='Табылмады'),404
    if kind=='users':
        if u.role!='admin' or x.id==u.id:return jsonify(error='Өшіруге болмайды'),403
    elif u.role!='admin' and x.class_name!=u.class_name:return jsonify(error='Рұқсат жоқ'),403
    db.session.delete(x); db.session.commit(); return jsonify(ok=True)

def seed():
    db.create_all()
    if User.query.count(): return
    admin_password=os.environ.get('ADMIN_PASSWORD','Admin2026!')
    admin=User(name='Әкімшілік',login='admin',password_hash=generate_password_hash(admin_password),role='admin',class_name='Барлық сынып',must_change=False)
    db.session.add(admin); db.session.commit()

with app.app_context(): seed()
if __name__=='__main__': app.run(host='0.0.0.0',port=int(os.environ.get('PORT','8000')))
